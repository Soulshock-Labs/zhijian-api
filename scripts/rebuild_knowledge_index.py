#!/usr/bin/env python3
"""
构建幼教知识库索引：
- 扫描 knowledge_base/incoming 与 knowledge_base/library
- 提取文本（docx/txt/md/json/csv）
- 关键词特征归类（国家标准/园本特色/教学模型/周计划等）
- 生成索引文件 knowledge_base/indexes/knowledge_index.json
- 生成画像路由结果 knowledge_base/indexes/profile_routes.json
"""

from __future__ import annotations

import csv
import hashlib
import json
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional at runtime until dependency is installed
    PdfReader = None

BASE_DIR = Path(__file__).resolve().parent.parent
KB_DIR = BASE_DIR / "knowledge_base"
SCAN_DIRS = (KB_DIR / "incoming", KB_DIR / "library")
INDEX_DIR = KB_DIR / "indexes"
PROFILE_FILE = KB_DIR / "profiles" / "profile_registry.json"
INDEX_FILE = INDEX_DIR / "knowledge_index.json"
ROUTE_FILE = INDEX_DIR / "profile_routes.json"

SUPPORTED_EXT = {".docx", ".txt", ".md", ".json", ".csv", ".pdf"}

FEATURE_KEYWORDS: dict[str, tuple[str, ...]] = {
    "national_standard": (
        "3-6岁", "3—6岁", "学习与发展指南", "五大领域", "健康", "语言", "社会", "科学", "艺术",
    ),
    "kindergarten_feature": (
        "园本特色", "本园特色", "课程特色", "主题活动", "特色活动", "节气", "非遗", "食育", "体智能",
    ),
    "teaching_model": (
        "蒙氏", "瑞吉欧", "华德福", "DAP", "项目化", "PBL", "游戏化", "自主游戏",
    ),
    "weekly_plan": (
        "周计划", "活动安排", "周目标", "周工作重点", "星期一", "星期二", "星期三", "星期四", "星期五",
    ),
    "daily_plan": (
        "日计划", "日教案", "导入", "过程", "延伸", "反思",
    ),
    "observation": (
        "观察记录", "观察要点", "评估", "评价", "记录建议",
    ),
    "safety_health": (
        "卫生保健", "安全教育", "生活活动", "如厕", "洗手", "离园", "午睡",
    ),
    "family_collab": (
        "家园共育", "家园配合", "家长工作", "亲子", "家庭任务",
    ),
    "environment_design": (
        "环境创设", "环创", "主题墙", "材料投放", "区域材料", "道具",
    ),
}

SOURCE_PRIORITY = {
    "national_standard": 300,
    "kindergarten_current": 200,
    "kindergarten_others": 120,
    "generic_library": 80,
    "incoming": 20,
}

PRIMARY_PRIORITY = (
    "weekly_plan",
    "daily_plan",
    "national_standard",
    "environment_design",
    "family_collab",
    "safety_health",
    "teaching_model",
    "kindergarten_feature",
    "observation",
)


@dataclass
class FileDoc:
    path: Path
    source: str
    text: str


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _safe_read_text(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=enc)
        except Exception:
            continue
    return ""


def _extract_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
    except Exception:
        return ""
    chunks: list[str] = []
    for p in doc.paragraphs:
        t = " ".join((p.text or "").split()).strip()
        if t:
            chunks.append(t)
    for table in doc.tables:
        for row in table.rows:
            line = " | ".join(" ".join((c.text or "").split()) for c in row.cells).strip(" |")
            if line:
                chunks.append(line)
    return "\n".join(chunks)


def _extract_json(path: Path) -> str:
    raw = _safe_read_text(path)
    if not raw:
        return ""
    try:
        obj = json.loads(raw)
        return json.dumps(obj, ensure_ascii=False)
    except Exception:
        return raw


def _extract_csv(path: Path) -> str:
    try:
        with path.open("r", encoding="utf-8") as f:
            rows = list(csv.reader(f))
    except Exception:
        return _safe_read_text(path)
    lines: list[str] = []
    for row in rows:
        line = " | ".join(x.strip() for x in row if x and x.strip())
        if line:
            lines.append(line)
    return "\n".join(lines)


def _extract_pdf(path: Path) -> str:
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(str(path))
    except Exception:
        return ""
    chunks: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = re.sub(r"\s+", " ", text).strip()
        if text:
            chunks.append(text)
    return "\n".join(chunks)


def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return _extract_docx(path)
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext in {".txt", ".md"}:
        return _safe_read_text(path)
    if ext == ".json":
        return _extract_json(path)
    if ext == ".csv":
        return _extract_csv(path)
    return ""


def _iter_docs() -> list[FileDoc]:
    docs: list[FileDoc] = []
    for scan_dir in SCAN_DIRS:
        if not scan_dir.exists():
            continue
        source = scan_dir.name
        for path in scan_dir.rglob("*"):
            if not path.is_file():
                continue
            if path.name.startswith("."):
                continue
            if path.suffix.lower() not in SUPPORTED_EXT:
                continue
            text = _extract_text(path).strip()
            docs.append(FileDoc(path=path, source=source, text=text))
    return docs


def _source_tier(path: Path, source: str) -> str:
    p = str(path).replace("\\", "/")
    if "/library/national_standard/" in p:
        return "national_standard"
    if "/library/kindergarten_standard/current/" in p:
        return "kindergarten_current"
    if "/library/kindergarten_standard/others/" in p:
        return "kindergarten_others"
    if source == "library":
        return "generic_library"
    return "incoming"


def _count_keywords(text: str) -> dict[str, int]:
    out: dict[str, int] = {}
    for feature, words in FEATURE_KEYWORDS.items():
        hits = 0
        for w in words:
            hits += text.count(w)
        out[feature] = hits
    return out


def _path_signal(path: Path) -> str:
    parts = [path.name]
    parts.extend(parent.name for parent in list(path.parents)[:4])
    return " ".join(parts)


def _tags_from_counts(counts: dict[str, int]) -> list[str]:
    return sorted([k for k, v in counts.items() if v > 0])


def _pick_primary(tags: list[str]) -> str:
    for key in PRIMARY_PRIORITY:
        if key in tags:
            return key
    return "mixed"


def _brief(text: str, max_len: int = 180) -> str:
    compact = re.sub(r"\s+", " ", text).strip()
    return compact[:max_len]


def _load_profiles() -> list[dict[str, Any]]:
    if not PROFILE_FILE.exists():
        return []
    try:
        data = json.loads(PROFILE_FILE.read_text(encoding="utf-8"))
        if isinstance(data, list):
            return [x for x in data if isinstance(x, dict)]
    except Exception:
        pass
    return []


def _score_for_profile(record: dict[str, Any], profile: dict[str, Any]) -> int:
    tags = set(record.get("tags", []))
    pref = set(profile.get("preferred_tags", []))
    avoid = set(profile.get("avoid_tags", []))
    score = len(tags & pref) * 10
    score -= len(tags & avoid) * 8
    score += int(record.get("feature_counts", {}).get("national_standard", 0))
    score += int(record.get("source_priority", 0))
    return score


def _build_routes(index_records: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    profiles = _load_profiles()
    out: dict[str, list[dict[str, Any]]] = {}
    for p in profiles:
        pid = str(p.get("id", "")).strip()
        if not pid:
            continue
        ranked = sorted(
            index_records,
            key=lambda rec: _score_for_profile(rec, p),
            reverse=True,
        )
        out[pid] = [
            {
                "doc_id": rec["doc_id"],
                "filename": rec["filename"],
                "path": rec["path"],
                "primary_bucket": rec["primary_bucket"],
                "tags": rec["tags"],
            }
            for rec in ranked[:50]
            if _score_for_profile(rec, p) > 0
        ]
    return out


def main() -> None:
    INDEX_DIR.mkdir(parents=True, exist_ok=True)
    docs = _iter_docs()
    records: list[dict[str, Any]] = []
    bucket_counts: dict[str, int] = {}

    for doc in docs:
        stat = doc.path.stat()
        text = doc.text
        signal_text = f"{text}\n{_path_signal(doc.path)}".strip()
        counts = _count_keywords(signal_text)
        tags = _tags_from_counts(counts)
        primary = _pick_primary(tags)
        tier = _source_tier(doc.path, doc.source)
        source_priority = SOURCE_PRIORITY.get(tier, 0)
        bucket_counts[primary] = bucket_counts.get(primary, 0) + 1
        digest = hashlib.sha1(str(doc.path).encode("utf-8")).hexdigest()[:12]

        records.append(
            {
                "doc_id": digest,
                "filename": doc.path.name,
                "path": str(doc.path),
                "source": doc.source,
                "source_tier": tier,
                "source_priority": source_priority,
                "size": stat.st_size,
                "mtime_utc": datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
                "primary_bucket": primary,
                "tags": tags,
                "feature_counts": counts,
                "excerpt": _brief(text) or _brief(_path_signal(doc.path)),
                "text_len": len(text),
            }
        )

    index_payload = {
        "version": "kb-v1",
        "generated_at_utc": _now_iso(),
        "doc_count": len(records),
        "buckets": bucket_counts,
        "records": records,
    }
    INDEX_FILE.write_text(json.dumps(index_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    routes = _build_routes(records)
    route_payload = {
        "version": "kb-route-v1",
        "generated_at_utc": _now_iso(),
        "profile_routes": routes,
    }
    ROUTE_FILE.write_text(json.dumps(route_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"[ok] indexed docs: {len(records)}")
    print(f"[ok] index file: {INDEX_FILE}")
    print(f"[ok] route file: {ROUTE_FILE}")
    print(f"[ok] buckets: {json.dumps(bucket_counts, ensure_ascii=False)}")


if __name__ == "__main__":
    main()
