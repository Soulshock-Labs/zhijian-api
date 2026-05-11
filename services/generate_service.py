from __future__ import annotations

import base64
import io
import json
import logging
import re
import textwrap
from uuid import uuid4

from docx import Document

from ai_service import _parse_json_payload, build_system_prompt
from core.settings import AI_MODEL, APP_VERSION, DASHSCOPE_API_KEY, _TEMP_EXPORT_DIR
from core.state import FIRESTORE_ENABLED, _TEMP_EXPORTS, _fs, client
from core.clients import _raise_if_invalid_dashscope_key
from core.utils import _utc_iso
from prompt_engineering.prompt_config import get_prompt_template
from services.data_store import _load_user_accounts
from word_engine.template_tools import _today_str

logger = logging.getLogger(__name__)

def _build_mini_doc_payload(
    filled_bytes: bytes,
    original_name: str,
    export_engine: str,
) -> dict[str, str]:
    token = uuid4().hex
    export_path = _TEMP_EXPORT_DIR / f"{token}.docx"
    export_path.write_bytes(filled_bytes)
    import time
    _TEMP_EXPORTS[token] = {
        "path": str(export_path),
        "filename": original_name,
        "engine": export_engine,
        "created_at": str(time.time()),
    }
    return {
        "status": "ok",
        "download_url": f"/mini-export/{token}",
        "filename": original_name,
        "engine": export_engine,
        "file_base64": base64.b64encode(filled_bytes).decode("ascii"),
    }

def _build_observation_prompt(
    theme: str,
    child_name: str,
    scene: str,
    note: str,
    phil: str,
    photo_names: list[str],
) -> str:
    photo_text = "、".join(photo_names) if photo_names else "（未上传照片）"
    child = child_name or "幼儿"
    return textwrap.dedent(
        f"""
        请为幼儿园教师生成一份「拍照观察记录」，只返回 JSON，不要其他文字。

        【输入信息】
        - 观察主题：{theme}
        - 观察对象：{child}
        - 观察场景：{scene}
        - 教育理念：{phil}
        - 照片文件名：{photo_text}
        - 教师补充：{note or "无"}

        【输出要求】
        - 语气专业、真实、可执行，避免空话
        - 贴近一线幼师记录语境
        - 每条建议可直接用于复盘和家园沟通

        【JSON 格式】
        {{
          "title": "观察记录标题",
          "summary": "观察概述（80-120字）",
          "records": ["关键观察1", "关键观察2", "关键观察3"],
          "analysis": "发展解读（80-120字）",
          "supports": ["支持策略1", "支持策略2", "支持策略3"],
          "home_cooperation": "家园共育建议（40-80字）",
          "next_plan": "下次跟进计划（40-80字）",
          "generated": "系统生成说明（简短）",
          "adjustment": "教师可调整建议（简短）"
        }}
        """
    ).strip()

def _mock_observation_content(
    theme: str,
    child_name: str,
    scene: str,
    note: str,
    photo_names: list[str],
) -> dict:
    child = child_name or "幼儿"
    topic = theme or "主题活动"
    note_text = f"教师补充：{note}" if note else "教师补充：无"
    return {
        "title": f"{topic}观察记录",
        "summary": (
            f"在{scene}中，{child}围绕「{topic}」表现出较高参与度，"
            "能够主动回应任务并与同伴互动。"
            "从现场照片与过程记录看，幼儿有持续投入与表达意愿。"
        ),
        "records": [
            f"{child}在活动中能主动操作材料，并保持阶段性专注。",
            "幼儿在同伴互动中出现协商与轮流行为，社交参与较积极。",
            "面对任务变化时，幼儿愿意尝试不同方法并表达自己的发现。",
        ],
        "analysis": (
            f"{child}在观察表达与合作参与方面已有可见进步。"
            "建议继续通过同类情境巩固经验，提升语言组织与问题解决深度。"
        ),
        "supports": [
            "提供可重复操作的低结构材料，支持幼儿再次探索。",
            "教师使用开放式提问，引导幼儿描述“如何发现”。",
            "设计双人协作小任务，强化沟通与角色分工体验。",
        ],
        "home_cooperation": "建议家长在家庭场景延续同主题观察活动，记录孩子的关键表达并反馈给教师。",
        "next_plan": "下次活动增加“分享与复述”环节，帮助幼儿梳理观察过程并形成表达闭环。",
        "generated": f"已接收 {len(photo_names)} 张照片；{note_text}",
        "adjustment": "教师可根据班级节奏对难度、材料与提问方式做个性化微调。",
    }

def generate_observation_content(
    theme: str,
    child_name: str,
    scene: str,
    note: str,
    phil: str,
    photo_names: list[str],
) -> dict:
    if not DASHSCOPE_API_KEY:
        return _mock_observation_content(theme, child_name, scene, note, photo_names)
    try:
        resp = client.chat.completions.create(
            model=AI_MODEL,
            messages=[
                {"role": "system", "content": build_system_prompt()},
                {
                    "role": "user",
                    "content": _build_observation_prompt(
                        theme=theme,
                        child_name=child_name,
                        scene=scene,
                        note=note,
                        phil=phil,
                        photo_names=photo_names,
                    ),
                },
            ],
            temperature=1,
            max_tokens=1200,
        )
        raw = resp.choices[0].message.content.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return json.loads(raw)
    except Exception as e:
        _raise_if_invalid_dashscope_key(e)
        return _mock_observation_content(theme, child_name, scene, note, photo_names)

def _build_observation_docx_bytes(
    content: dict,
    theme: str,
    child_name: str,
    scene: str,
    phil: str,
    note: str,
    photo_names: list[str],
) -> bytes:
    doc = Document()
    today = _today_str()
    child = child_name or "幼儿"
    doc.add_heading(content.get("title", f"{theme}观察记录"), level=1)

    intro = doc.add_paragraph()
    intro.add_run("日期：").bold = True
    intro.add_run(f"{today}    ")
    intro.add_run("观察场景：").bold = True
    intro.add_run(f"{scene}    ")
    intro.add_run("观察对象：").bold = True
    intro.add_run(child)

    doc.add_paragraph(f"教育理念：{phil}")
    doc.add_paragraph(f"照片数量：{len(photo_names)}")
    if photo_names:
        doc.add_paragraph("照片文件：")
        for name in photo_names:
            doc.add_paragraph(name, style="List Bullet")

    if note:
        doc.add_heading("教师补充", level=2)
        doc.add_paragraph(note)

    doc.add_heading("一、观察概述", level=2)
    doc.add_paragraph(content.get("summary", "（待补充）"))

    doc.add_heading("二、关键观察记录", level=2)
    records = content.get("records", []) or []
    if records:
        for row in records:
            doc.add_paragraph(str(row), style="List Bullet")
    else:
        doc.add_paragraph("（待补充）")

    doc.add_heading("三、发展解读", level=2)
    doc.add_paragraph(content.get("analysis", "（待补充）"))

    doc.add_heading("四、支持策略", level=2)
    supports = content.get("supports", []) or []
    if supports:
        for row in supports:
            doc.add_paragraph(str(row), style="List Bullet")
    else:
        doc.add_paragraph("（待补充）")

    doc.add_heading("五、家园共育建议", level=2)
    doc.add_paragraph(content.get("home_cooperation", "（待补充）"))

    doc.add_heading("六、下次跟进计划", level=2)
    doc.add_paragraph(content.get("next_plan", "（待补充）"))

    doc.add_heading("七、生成", level=2)
    doc.add_paragraph(content.get("generated", "由系统生成初稿"))

    doc.add_heading("八、调整", level=2)
    doc.add_paragraph(content.get("adjustment", "教师可在此补充个体化调整建议"))

    try:
        doc.core_properties.comments = (
            f"小纸笺导出 · v{APP_VERSION} · 排版引擎 python-docx（观察记录）"
        )
    except Exception:
        pass
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()

async def _generate_weekly_for_user(
    openid: str,
    theme: str,
    phil: str,
    activities: list[str],
    class_level: str,
) -> dict:
    """单用户异步生成周计划，注入该用户的 agent_profile。"""
    try:
        accounts = _load_user_accounts()
        entry = accounts.get(openid, {})
        agent = entry.get("agent_profile", {})

        # 把 agent 性格注入到 prompt 里
        prompt_template = get_prompt_template()
        system_prompt = prompt_template.build_system_prompt()
        if agent:
            agent_hint = (
                f"\n\n【本次老师的智能体设定】\n"
                f"名字：{agent.get('name','小助手')}\n"
                f"性格：{agent.get('personality','热心、耐心')}\n"
                f"音调：{agent.get('tone','亲切温暖')}\n"
                f"教学风格：{agent.get('style','鼓励式教学')}\n"
                f"请按照以上风格生成内容，让老师感受到专属感。"
            )
            system_prompt = system_prompt + agent_hint

        user_prompt = prompt_template.build_user_prompt(
            theme=theme,
            class_level=class_level or "中班",
            philosophy=phil,
            activities=activities or ["区域活动", "户外活动"],
        )

        resp = client.chat.completions.create(
            model=AI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=1,
            max_tokens=4096,
        )
        raw = resp.choices[0].message.content.strip()
        plan = _parse_json_payload(raw)

        # 写入 Firestore 历史
        if FIRESTORE_ENABLED:
            try:
                _fs().collection("users").document(openid)\
                    .collection("history").add({
                        "type": "weekly_plan",
                        "theme": theme,
                        "plan": plan,
                        "created_at": _utc_iso(),
                    })
            except Exception as e:
                logger.warning("Firestore 写历史失败：%s", e)

        return {"openid": openid, "ok": True, "plan": plan}
    except Exception as e:
        logger.error("批量生成失败 openid=%s err=%s", openid, e)
        return {"openid": openid, "ok": False, "error": str(e)}
