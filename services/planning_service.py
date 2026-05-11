"""
services/planning_service.py — 周日计划业务逻辑
职责：纯业务函数，不依赖 FastAPI，不处理 HTTP。
路由层调用这里的函数，不在这里处理 Request/Response。

包含：
· 周计划 AI 生成（generate_weekly_content）
· 日教案 AI 生成（generate_daily_content）
· 学期骨架生成（build_term_month_week_skeleton）
· 日教案 Word 填充（fill_daily_word_template）
· 日教案关键字映射（DAILY_CELL_KEYWORD_MAP / match_daily_field）
"""
from __future__ import annotations

import io
import json
import logging
import re
import textwrap
from typing import Optional

from docx import Document
from fastapi import HTTPException

from core.settings import AI_MODEL, AI_MODEL_FAST, APP_VERSION, DASHSCOPE_API_KEY
from core.state import client, deepseek_client, qwen_client, ENABLE_ASPOSE_WORDS
from core.clients import _raise_if_invalid_dashscope_key


def _resolve_client(model_name: str):
    """根据模型名称前缀选择对应的 API client。"""
    if model_name.startswith("deepseek"):
        if deepseek_client is None:
            raise HTTPException(status_code=503, detail="DeepSeek API 未配置（缺少 DEEPSEEK_API_KEY）")
        return deepseek_client
    if model_name.startswith("qwen"):
        if qwen_client is None:
            raise HTTPException(status_code=503, detail="Qwen API 未配置（缺少 QWEN_API_KEY）")
        return qwen_client
    # 默认走 Moonshot（当前 client）
    return client
from word_engine.field_map import PHILOSOPHY_HINTS
from word_engine.template_tools import _today_str, _get_cell_text, _is_colored_cell, clean_template_keep_style
from word_engine.docx_filler import _write_cell_preserve_style
from word_engine.aspose_filler import (
    _aw_require,
    _aw_doc_from_bytes,
    _aw_doc_to_bytes,
    _aw_cell_text,
    _aw_cell_has_color,
    _aw_cell_dedup_key,
    _aw_write_cell_preserve_style,
    _aw_stamp_export_provenance,
    _aw_node_as_table,
)
from ai_service import build_system_prompt
from prompt_engineering.prompt_config import get_prompt_template

logger = logging.getLogger(__name__)

# ══════════════════════════════════════════════════════════════════════
# 日教案单元格关键字映射
# ══════════════════════════════════════════════════════════════════════

DAILY_CELL_KEYWORD_MAP: list[tuple[list[str], str]] = [
    (["活动名称", "课题名称", "活动标题", "课题"],    "title"),
    (["教育理念", "课程理念", "理念"],                "philosophy"),
    (["班级", "年龄段", "年龄", "班"],                "class_info"),
    (["活动目标", "教学目标", "目标"],                "goals"),
    (["活动准备", "材料准备", "准备"],                "preparation"),
    (["活动导入", "情境导入", "导入", "引入"],         "introduction"),
    (["活动过程", "基本过程", "主要过程", "过程"],     "process"),
    (["活动延伸", "延伸活动", "延伸", "拓展"],        "extension"),
    (["活动反思", "教师反思", "反思", "评价小结"],     "reflection"),
    (["观察要点", "观察记录", "观察重点"],             "observation"),
    (["日期", "时间"],                                "date_info"),
]


def match_daily_field(cell_text: str) -> Optional[str]:
    """将单元格文字映射到日教案字段名，未命中返回 None。"""
    t = cell_text.strip()
    for keywords, field in DAILY_CELL_KEYWORD_MAP:
        for kw in keywords:
            if kw in t:
                return field
    return None


# ══════════════════════════════════════════════════════════════════════
# 周计划 AI 生成
# ══════════════════════════════════════════════════════════════════════

def build_weekly_prompt(
    theme: str,
    phil: str,
    activities: list[str],
    class_level: str = "",
) -> str:
    """使用 Prompt 工程系统生成周计划 User Prompt（可复现）。"""
    prompt_template = get_prompt_template()
    return prompt_template.build_user_prompt(
        theme=theme,
        class_level=class_level or "中班",
        philosophy=phil,
        activities=activities or ["区域活动", "户外活动"],
    )


def _analyze_doc_with_deepseek(doc_md: str, theme: str, phil: str) -> str:
    """
    用 DeepSeek 分析老师上传的文档 Markdown，提取风格/偏好/结构要点。
    返回一段精炼的「文档分析摘要」字符串，用于注入周计划生成 prompt。
    若 DeepSeek 不可用则直接返回原始 doc_md（降级处理）。
    """
    from core.state import deepseek_client
    from core.settings import DEEPSEEK_API_KEY

    if not DEEPSEEK_API_KEY or deepseek_client is None:
        logger.warning("DeepSeek 未配置，跳过文档分析，直接使用原文摘要")
        return doc_md

    analyze_prompt = f"""你是一名幼儿园课程设计专家。老师上传了一份参考文档，请你仔细阅读并提炼以下信息：

1. **文档类型**：周计划 / 日教案 / 活动方案 / 其他
2. **周主题/活动主题**：从文档中提取
3. **教育理念风格**：老师倾向的教学方式和理念关键词
4. **内容结构特点**：文档的组织方式、常用板块
5. **活动类型偏好**：频繁出现的活动类型（区域、户外、艺术等）
6. **写作风格**：语言风格（简洁/详细、学术/生活化等）
7. **可复用的亮点内容**：值得在新计划中延续或参考的具体做法

---

【本次生成目标】
- 新周主题：{theme}
- 教育理念：{phil}

请基于文档内容，给出5-8条具体建议，帮助生成与老师风格一致的新周计划。
用简洁的要点形式输出，不需要 JSON，直接输出中文要点列表。

---

【老师上传的文档内容】
{doc_md}
"""

    try:
        resp = deepseek_client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "你是专业的幼儿园课程顾问，擅长分析教师文档并提炼风格特征。"},
                {"role": "user", "content": analyze_prompt},
            ],
            temperature=0.3,
            max_tokens=1000,
        )
        analysis = resp.choices[0].message.content.strip()
        logger.info("DeepSeek 文档分析完成，摘要长度：%d", len(analysis))
        return analysis
    except Exception as e:
        logger.warning("DeepSeek 文档分析失败，降级使用原文摘要：%s", e)
        return doc_md


def generate_weekly_content(
    theme: str,
    phil: str,
    activities: list[str],
    class_level: str = "",
    model: str = "",
    doc_md: str = "",
) -> dict:
    """调用 AI 生成五天周计划 JSON，未配置 Key 时返回 Mock 数据。
    model: 可指定模型，空字符串则使用 AI_MODEL_FAST。
    doc_md: 老师上传文档的 Markdown 内容，非空时先用 DeepSeek 分析再注入 prompt。
    """
    if not DASHSCOPE_API_KEY:
        return _mock_weekly(theme, phil)
    model_to_use = model.strip() if model.strip() else AI_MODEL_FAST

    # ── 文档分析：若老师上传了参考文档，先用 DeepSeek 分析提炼风格 ──
    doc_analysis_hint = ""
    if doc_md.strip():
        analysis = _analyze_doc_with_deepseek(doc_md.strip(), theme, phil)
        if analysis.strip():
            doc_analysis_hint = (
                "\n\n【老师上传的参考文档分析】\n"
                "以下是对老师上传文档的分析，请参考这些风格特点生成本次周计划：\n"
                f"{analysis}\n"
                "【请在保持以上风格的基础上，结合本次主题和理念生成新的周计划】"
            )
            logger.info("文档分析摘要已注入 prompt，长度：%d", len(doc_analysis_hint))

    def _build_messages() -> list[dict]:
        prompt_template = get_prompt_template()
        system_content = prompt_template.build_system_prompt() + doc_analysis_hint
        user_content = build_weekly_prompt(theme, phil, activities, class_level)
        return [
            {"role": "system", "content": system_content},
            {"role": "user",   "content": user_content},
        ]

    try:
        resp = _resolve_client(model_to_use).chat.completions.create(
            model=model_to_use,
            messages=_build_messages(),
            temperature=1,
            max_tokens=4096,
        )
        raw = resp.choices[0].message.content.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return json.loads(raw)
    except Exception as e:
        import traceback, time
        logger.error("generate_weekly_content attempt failed: %s\n%s", e, traceback.format_exc())
        _raise_if_invalid_dashscope_key(e)
        # 超时或临时错误：重试一次
        try:
            time.sleep(2)
            resp = _resolve_client(model_to_use).chat.completions.create(
                model=model_to_use,
                messages=_build_messages(),
                temperature=1,
                max_tokens=4096,
            )
            raw = resp.choices[0].message.content.strip()
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)
            return json.loads(raw)
        except Exception as e2:
            logger.error("generate_weekly_content retry also failed: %s", e2)
            raise HTTPException(status_code=502, detail=f"周计划生成失败（已重试）：{e2}")


def _mock_weekly(theme: str, phil: str) -> dict:
    tasks = [
        ("观察与记录", "感官探索", "区域活动"),
        ("创意表达",   "艺术创作", "环创活动"),
        ("合作探究",   "社交能力", "户外活动"),
        ("生活技能",   "自理能力", "生活活动"),
        ("成果分享",   "语言表达", "家园活动"),
    ]
    return {
        "week_theme": theme,
        "philosophy": phil,
        "days": [
            {
                "day":           f"周{c}",
                "task":          f"{theme}·{t}",
                "focus":         f,
                "activity_type": at,
                "hint":          f"关注幼儿在{f}方面的表现",
            }
            for (t, f, at), c in zip(tasks, ["一", "二", "三", "四", "五"])
        ],
    }


# ══════════════════════════════════════════════════════════════════════
# 学期骨架生成（园部 To B）
# ══════════════════════════════════════════════════════════════════════

def build_term_month_week_skeleton(
    term_theme: str,
    start_month: int,
    month_count: int,
) -> dict:
    """生成园部学期→月→周三级骨架，供园所统一规划后下钻班级执行。"""
    month_count = max(1, min(month_count, 6))
    months: list[dict] = []
    for i in range(month_count):
        month_no    = ((start_month - 1 + i) % 12) + 1
        month_title = f"{month_no}月"
        weeks = [
            {
                "week_index": w + 1,
                "week_theme": f"{term_theme}·{month_title}第{w + 1}周",
                "focus":      "园所统一目标待补充",
                "status":     "planned",
            }
            for w in range(4)
        ]
        months.append({
            "month":      month_title,
            "month_goal": "围绕学期目标分解月度重点",
            "weeks":      weeks,
        })
    return {"term_theme": term_theme, "months": months}


# ══════════════════════════════════════════════════════════════════════
# 日教案 AI 生成
# ══════════════════════════════════════════════════════════════════════

def build_daily_prompt(
    week_theme: str,
    day: str,
    task: str,
    phil: str,
    phil_hint: str,
) -> str:
    return textwrap.dedent(f"""
        请为幼儿园日教案生成专业内容，以 JSON 格式返回。

        【基本信息】
        - 周主题：{week_theme}
        - 今日活动（{day}）：{task}
        - 教育理念：{phil}

        【理念专业词汇要求】
        {phil_hint}

        【四维结构要求】
        请严格按照「导入 → 过程 → 延伸 → 反思」四个维度设计，每个维度字数控制在100-150字。

        【JSON 输出格式（只返回 JSON）】
        {{
          "title": "活动名称（即今日任务）",
          "goals": [
            "目标1（含维度前缀，如感知/语言/社会等）",
            "目标2",
            "目标3"
          ],
          "preparation": ["材料1", "材料2", "材料3"],
          "introduction": "导入环节（5-8分钟）：情境创设、激发兴趣、连接已有经验的具体步骤",
          "process": "活动过程（20-25分钟）：分步骤描述教师引导动作、幼儿操作内容、关键问题设计",
          "extension": "延伸活动（5-10分钟）：区域延伸、家园延伸或跨日连接建议",
          "reflection": "教师反思：今日活动的观察重点与课后反思问题（含{phil}理念专业术语）",
          "observation": "重点观察要点（3条，每条15字以内）"
        }}
    """).strip()


def _extract_json_block(text: str) -> str:
    """从可能包含额外说明的文本中提取最外层 JSON 对象。"""
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```\s*$", "", text, flags=re.IGNORECASE)
    # 如果仍有前后缀，尝试定位第一个 { 和最后一个配对的 }
    start = text.find("{")
    if start == -1:
        return text
    depth = 0
    end = -1
    for i, ch in enumerate(text[start:], start=start):
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                end = i
                break
    if end != -1:
        return text[start:end + 1]
    return text


def generate_daily_content(
    week_theme: str,
    day: str,
    task: str,
    phil: str,
) -> dict:
    """调用 AI 生成四维日教案 JSON，未配置 Key 时返回 Mock 数据。"""
    if not DASHSCOPE_API_KEY:
        return _mock_daily(week_theme, day, task, phil)
    phil_hint = PHILOSOPHY_HINTS.get(phil, "")
    try:
        resp = client.chat.completions.create(
            model=AI_MODEL,
            messages=[
                {"role": "system", "content": build_system_prompt()},
                {"role": "user",   "content": build_daily_prompt(week_theme, day, task, phil, phil_hint)},
            ],
            temperature=1,
            max_tokens=4096,
        )
        raw = resp.choices[0].message.content.strip()
        raw = _extract_json_block(raw)
        return json.loads(raw)
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=502, detail=f"日教案内容解析失败：{e}")
    except Exception as e:
        _raise_if_invalid_dashscope_key(e)
        raise HTTPException(status_code=502, detail=f"日教案生成失败：{e}")


def _mock_daily(week_theme: str, day: str, task: str, phil: str) -> dict:
    return {
        "title": task,
        "goals": [
            f"感知「{task}」的核心特征，积累相关直接经验",
            "通过操作与探索，发展观察比较能力",
            "在互动中提升语言表达与合作意识",
        ],
        "preparation": [f"「{task}」相关操作材料", "观察记录单", "展示板"],
        "introduction": (
            f"（{day}·导入）以情境导入：展示与「{task}」相关的真实物品或图片，"
            f"提问：'你在哪里见过这个？它让你想到了什么？'激活幼儿先备经验，"
            f"自然过渡到今日探究活动。（{phil}理念指导：注重儿童已有经验的联结）"
        ),
        "process": (
            f"（过程）①教师示范操作步骤，明确今日探究任务；"
            f"②幼儿自主操作，教师巡回观察并记录关键行为；"
            f"③小组分享：'你发现了什么？你是怎么做到的？'；"
            f"④集体总结，提炼「{task}」的核心经验。"
        ),
        "extension": (
            f"（延伸）区域延伸：在科学探究区/美工区延续「{task}」相关材料供幼儿自主探索；"
            "家园延伸：请家长与孩子在家寻找与「{task}」相关的生活场景并拍照分享。"
        ),
        "reflection": (
            f"（{phil}·反思）观察今日幼儿在「{task}」活动中的参与度与深度思考迹象；"
            "记录有价值的幼儿语言；思考：哪些材料激发了更持久的探究？"
            "明日如何在此基础上推进？"
        ),
        "observation": "观察专注投入时长 | 记录语言表达关键词 | 关注合作协商行为",
    }


# ══════════════════════════════════════════════════════════════════════
# 日教案 Word 填充
# ══════════════════════════════════════════════════════════════════════

def _build_daily_fill_data(
    daily_content: dict,
    week_theme: str,
    day: str,
    phil: str,
) -> dict[str, str]:
    today = _today_str()
    goals_text = "\n".join(
        f"目标{i+1}：{g}" for i, g in enumerate(daily_content.get("goals", []))
    )
    prep_text = "\n".join(f"• {p}" for p in daily_content.get("preparation", []))
    obs_raw = daily_content.get("observation", "")
    if isinstance(obs_raw, list):
        obs_items = [str(o).strip() for o in obs_raw if str(o).strip()]
    else:
        obs_raw_text = str(obs_raw)
        obs_items = [
            o.strip()
            for o in (obs_raw_text.split("|") if "|" in obs_raw_text else [obs_raw_text])
            if o.strip()
        ]
    obs_text = "\n".join(f"△ {o}" for o in obs_items)
    return {
        "title":        daily_content.get("title", f"{week_theme}·{day}"),
        "philosophy":   phil,
        "class_info":   f"________班    日期：{today}",
        "goals":        goals_text,
        "preparation":  prep_text,
        "introduction": daily_content.get("introduction", ""),
        "process":      daily_content.get("process", ""),
        "extension":    daily_content.get("extension", ""),
        "reflection":   daily_content.get("reflection", ""),
        "observation":  obs_text,
        "date_info":    f"{today}（{day}）",
    }


def _build_daily_structured_docx_bytes(
    daily_content: dict,
    week_theme: str,
    day: str,
    phil: str,
    source_day: Optional[dict] = None,
) -> bytes:
    """生成非表格结构化日计划 Word 文档（无模板时使用）。"""
    doc   = Document()
    today = _today_str()
    title = daily_content.get("title", f"{week_theme}·{day}日计划")
    doc.add_heading(f"{title}（{day}）", level=1)

    intro = doc.add_paragraph()
    intro.add_run("周主题：").bold = True
    intro.add_run(f"{week_theme}    ")
    intro.add_run("日期：").bold = True
    intro.add_run(f"{today}    ")
    intro.add_run("教育理念：").bold = True
    intro.add_run(phil)

    goals   = daily_content.get("goals", []) or []
    prep    = daily_content.get("preparation", []) or []
    obs_raw = daily_content.get("observation", "") or ""
    if isinstance(obs_raw, list):
        obs_items = [str(x).strip() for x in obs_raw if str(x).strip()]
    else:
        obs_text  = str(obs_raw)
        obs_items = [x.strip() for x in (obs_text.split("|") if "|" in obs_text else [obs_text]) if x.strip()]

    doc.add_heading("一、活动目标", level=2)
    for g in goals:
        doc.add_paragraph(str(g), style="List Bullet")
    if not goals:
        doc.add_paragraph("（待补充）")

    doc.add_heading("二、活动准备", level=2)
    for p in prep:
        doc.add_paragraph(str(p), style="List Bullet")
    if not prep:
        doc.add_paragraph("（待补充）")

    doc.add_heading("三、活动导入", level=2)
    doc.add_paragraph(daily_content.get("introduction", "（待补充）"))

    doc.add_heading("四、活动过程", level=2)
    doc.add_paragraph(daily_content.get("process", "（待补充）"))

    doc.add_heading("五、活动延伸", level=2)
    doc.add_paragraph(daily_content.get("extension", "（待补充）"))

    doc.add_heading("六、活动反思", level=2)
    doc.add_paragraph(daily_content.get("reflection", "（待补充）"))

    doc.add_heading("七、重点观察", level=2)
    for o in obs_items:
        doc.add_paragraph(o, style="List Bullet")
    if not obs_items:
        doc.add_paragraph("（待补充）")

    doc.add_heading("八、生成", level=2)
    gen = doc.add_paragraph()
    gen.add_run("来源：").bold = True
    gen.add_run("周计划联动自动生成\n")
    gen.add_run("模型：").bold = True
    gen.add_run(f"{AI_MODEL}\n")
    gen.add_run("生成时间：").bold = True
    gen.add_run(f"{today}\n")
    if source_day:
        doc.add_paragraph(
            f"任务：{source_day.get('task', '')}；"
            f"关注点：{source_day.get('focus', '')}；"
            f"活动类型：{source_day.get('activity_type', '')}；"
            f"提示：{source_day.get('hint', '')}"
        )

    doc.add_heading("九、调整", level=2)
    doc.add_paragraph("教师二次调整记录：")
    doc.add_paragraph("1. 今日微调点：________________________________________")
    doc.add_paragraph("2. 幼儿响应观察：______________________________________")
    doc.add_paragraph("3. 明日延伸计划：______________________________________")

    try:
        doc.core_properties.comments = (
            f"小纸笺导出 · v{APP_VERSION} · 排版引擎 python-docx（结构化日计划）"
        )
    except Exception:
        pass
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _fill_daily_template_docx_bytes(
    cleaned_bytes: bytes,
    fill_data: dict[str, str],
) -> bytes:
    """python-docx 日教案模板填充（Aspose 回退版）。"""
    doc            = Document(io.BytesIO(cleaned_bytes))
    filled_tc_elems: set = set()
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            cells = row.cells
            for col_idx, cell in enumerate(cells):
                field = match_daily_field(_get_cell_text(cell))
                if not field:
                    continue
                content = fill_data.get(field)
                if not content:
                    continue
                target = None
                if col_idx + 1 < len(cells):
                    cand = cells[col_idx + 1]
                    if not match_daily_field(_get_cell_text(cand)) and not _is_colored_cell(cand):
                        target = cand
                if (target is None
                        and row_idx + 1 < len(table.rows)
                        and col_idx < len(table.rows[row_idx + 1].cells)):
                    cand = table.rows[row_idx + 1].cells[col_idx]
                    if not match_daily_field(_get_cell_text(cand)) and not _is_colored_cell(cand):
                        target = cand
                if target is None:
                    continue
                tc_elem = target._tc
                if tc_elem in filled_tc_elems:
                    continue
                filled_tc_elems.add(tc_elem)
                _write_cell_preserve_style(target, content)
    try:
        doc.core_properties.comments = (
            f"小纸笺导出 · v{APP_VERSION} · 排版引擎 python-docx（Aspose 回退）"
        )
    except Exception:
        pass
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _fill_daily_template_aspose_bytes(
    cleaned_bytes: bytes,
    fill_data: dict[str, str],
) -> bytes:
    """Aspose.Words 日教案模板填充（保留原样式）。"""
    _aw_require()
    import aspose.words as aw  # noqa: PLC0415 — 懒加载
    doc          = _aw_doc_from_bytes(cleaned_bytes)
    filled_cells: set[int] = set()
    tables       = doc.get_child_nodes(aw.NodeType.TABLE, True)
    for ti in range(tables.count):
        table = _aw_node_as_table(tables[ti])
        for row_idx in range(table.rows.count):
            row = table.rows[row_idx]
            for col_idx in range(row.cells.count):
                cell  = row.cells[col_idx]
                field = match_daily_field(_aw_cell_text(cell))
                if not field:
                    continue
                content = fill_data.get(field)
                if not content:
                    continue
                target = None
                if col_idx + 1 < row.cells.count:
                    cand = row.cells[col_idx + 1]
                    if not match_daily_field(_aw_cell_text(cand)) and not _aw_cell_has_color(cand):
                        target = cand
                if (target is None
                        and row_idx + 1 < table.rows.count
                        and col_idx < table.rows[row_idx + 1].cells.count):
                    cand = table.rows[row_idx + 1].cells[col_idx]
                    if not match_daily_field(_aw_cell_text(cand)) and not _aw_cell_has_color(cand):
                        target = cand
                if target is None:
                    continue
                key = _aw_cell_dedup_key(target)
                if key in filled_cells:
                    continue
                filled_cells.add(key)
                _aw_write_cell_preserve_style(doc, target, content)
    _aw_stamp_export_provenance(doc)
    return _aw_doc_to_bytes(doc)


def fill_daily_word_template(
    template_bytes: bytes,
    daily_content: dict,
    week_theme: str,
    day: str,
    phil: str,
) -> tuple[bytes, str]:
    """
    日教案 Word 填充入口：Aspose 优先，失败自动回退 python-docx。
    返回 (docx 字节, 引擎标识)。
    """
    cleaned_bytes = clean_template_keep_style(template_bytes)
    fill_data     = _build_daily_fill_data(daily_content, week_theme, day, phil)
    try:
        if ENABLE_ASPOSE_WORDS:
            return _fill_daily_template_aspose_bytes(cleaned_bytes, fill_data), "aspose-words"
    except Exception as e:
        logger.warning("Aspose 日教案填充失败，回退 python-docx：%s", e, exc_info=True)
    return _fill_daily_template_docx_bytes(cleaned_bytes, fill_data), "python-docx"
