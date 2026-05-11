"""
word_engine/doc_reader.py — 老师上传文档理解引擎

职责：
  从老师上传的 .docx / .pdf / 图片 中提取结构化文本，
  生成可直接注入 prompt 的「参考文档摘要」。

核心函数：
  extract_doc_context(file_bytes, filename) -> DocContext
  format_for_prompt(ctx) -> str      # 注入 system/user prompt 的字符串
  to_markdown(ctx) -> str            # 转结构化 Markdown
  extract_image_context(file_bytes, filename) -> DocContext  # 图片 OCR（Kimi Vision）

设计原则：
  - 只提取文字，不依赖布局/样式（跨格式统一）
  - 表格优先：幼师文档大多数内容在表格里
  - 截断保护：单文档最多注入 3000 字，避免超 token
  - 失败静默：提取出错时返回空摘要，不中断生成流程
  - 图片走 Kimi moonshot-v1-8k-vision-preview，base64 直传
"""
from __future__ import annotations

import base64
import io
import logging
import re
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)

MAX_CHARS = 3000  # 注入 prompt 的最大字符数


# ──────────────────────────────────────────────
# 数据结构
# ──────────────────────────────────────────────

@dataclass
class DocContext:
    """从老师文档提取的结构化上下文。"""
    filename: str = ""
    file_type: str = ""          # "docx" | "pdf" | "unknown"
    raw_text: str = ""           # 全量提取文本（截断前）
    tables: list[list[list[str]]] = field(default_factory=list)  # [table][row][cell]
    summary: str = ""            # 截断后的摘要，用于 prompt 注入
    char_count: int = 0
    table_count: int = 0
    ok: bool = True
    error: str = ""


# ──────────────────────────────────────────────
# 主入口
# ──────────────────────────────────────────────

IMAGE_EXTS = {"jpg", "jpeg", "png", "webp", "gif"}

def extract_doc_context(file_bytes: bytes, filename: str) -> DocContext:
    """
    从上传文件提取文档上下文。
    支持 .docx / .pdf / 图片（jpg/jpeg/png/webp/gif）。
    图片走 Kimi Vision OCR，失败时返回 ok=False 的空 DocContext。
    """
    ctx = DocContext(filename=filename)
    ext = (filename or "").lower().rsplit(".", 1)[-1]

    if ext in {"docx", "pdf"}:
        ctx.file_type = ext
    elif ext in IMAGE_EXTS:
        ctx.file_type = "image"
    else:
        ctx.file_type = "unknown"

    try:
        if ctx.file_type == "docx":
            _extract_docx(file_bytes, ctx)
        elif ctx.file_type == "pdf":
            _extract_pdf(file_bytes, ctx)
        elif ctx.file_type == "image":
            _extract_image(file_bytes, ctx)
        else:
            ctx.ok = False
            ctx.error = f"不支持的文件格式：{ext}，支持 .docx、.pdf、.jpg、.jpeg、.png、.webp"
            return ctx

        ctx.summary = _build_summary(ctx)
        ctx.char_count = len(ctx.summary)
    except Exception as e:
        logger.warning("文档提取失败 [%s]: %s", filename, e, exc_info=True)
        ctx.ok = False
        ctx.error = str(e)

    return ctx


def format_for_prompt(ctx: DocContext) -> str:
    """
    将 DocContext 格式化为可直接拼入 prompt 的字符串。
    若提取失败或为空，返回空字符串（不影响生成）。
    """
    if not ctx.ok or not ctx.summary.strip():
        return ""

    return (
        f"【老师上传的参考文档：{ctx.filename}】\n"
        f"{ctx.summary}\n"
        f"【参考文档结束，请基于以上内容理解老师的风格和要求】"
    )


def to_markdown(ctx: DocContext) -> str:
    """
    将 DocContext 转换为结构化 Markdown 字符串，供 DeepSeek 分析。
    表格用 Markdown 表格格式输出，段落保留原文。
    """
    if not ctx.ok or not ctx.raw_text.strip():
        return ""

    parts: list[str] = [f"# 参考文档：{ctx.filename}\n"]

    # 如果有表格，优先输出 Markdown 表格格式
    if ctx.tables:
        parts.append("## 文档表格内容\n")
        for t_idx, table in enumerate(ctx.tables, 1):
            if not table:
                continue
            parts.append(f"### 表格 {t_idx}\n")
            # 第一行作表头
            header = table[0]
            parts.append("| " + " | ".join(header) + " |")
            parts.append("| " + " | ".join(["---"] * len(header)) + " |")
            for row in table[1:]:
                # 对齐列数
                while len(row) < len(header):
                    row.append("")
                parts.append("| " + " | ".join(row[:len(header)]) + " |")
            parts.append("")

    # 段落文本（去掉已在表格里的内容，减少重复）
    if ctx.raw_text.strip():
        parts.append("## 文档正文内容\n")
        # 截断到 MAX_CHARS 避免超 token
        text = ctx.raw_text
        if len(text) > MAX_CHARS:
            text = text[:MAX_CHARS] + f"\n\n…（已截取，原文共 {len(ctx.raw_text)} 字）"
        parts.append(text)

    return "\n".join(parts)


# ──────────────────────────────────────────────
# .docx 提取
# ──────────────────────────────────────────────

def _extract_docx(file_bytes: bytes, ctx: DocContext) -> None:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    # 1. 段落文本
    para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if para_texts:
        parts.append("\n".join(para_texts))

    # 2. 表格（幼师文档核心内容区）
    for table in doc.tables:
        ctx.table_count += 1
        table_rows: list[list[str]] = []
        row_texts: list[str] = []

        for row in table.rows:
            cells = [_clean_cell(c.text) for c in row.cells]
            # 合并单元格会重复，去重相邻
            deduped = _dedup_adjacent(cells)
            non_empty = [c for c in deduped if c]
            if non_empty:
                row_texts.append(" | ".join(non_empty))
                table_rows.append(non_empty)

        if row_texts:
            parts.append("\n".join(row_texts))
            ctx.tables.append(table_rows)

    ctx.raw_text = "\n\n".join(parts)


# ──────────────────────────────────────────────
# .pdf 提取
# ──────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes, ctx: DocContext) -> None:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise RuntimeError("pypdf 未安装，无法读取 PDF。请运行 pip install pypdf")

    reader = PdfReader(io.BytesIO(file_bytes))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            pages.append(text)

    ctx.raw_text = "\n\n".join(pages)


# ──────────────────────────────────────────────
# 摘要构建（截断到 MAX_CHARS）
# ──────────────────────────────────────────────

def _build_summary(ctx: DocContext) -> str:
    text = ctx.raw_text

    # 清理多余空白
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = text.strip()

    if len(text) <= MAX_CHARS:
        return text

    # 截断时保留开头（通常是主题/目标）和结尾（通常是反思/要求）
    head = text[: MAX_CHARS * 2 // 3]
    tail = text[-(MAX_CHARS // 3):]

    # 在句子边界截断
    head = _cut_at_sentence(head)
    tail = _cut_at_sentence_start(tail)

    return head + f"\n\n…（文档较长，已截取关键段落，共 {len(text)} 字）…\n\n" + tail


def _cut_at_sentence(text: str) -> str:
    """在最后一个句子边界截断。"""
    for sep in ("。", "；", "\n"):
        idx = text.rfind(sep)
        if idx > len(text) * 0.7:
            return text[: idx + 1]
    return text


def _cut_at_sentence_start(text: str) -> str:
    """从第一个句子边界开始。"""
    for sep in ("。", "；", "\n"):
        idx = text.find(sep)
        if 0 < idx < len(text) * 0.3:
            return text[idx + 1 :]
    return text


# ──────────────────────────────────────────────
# 图片 OCR（Kimi Vision）
# ──────────────────────────────────────────────

# 图片 MIME 类型映射
_IMAGE_MIME: dict[str, str] = {
    "jpg":  "image/jpeg",
    "jpeg": "image/jpeg",
    "png":  "image/png",
    "webp": "image/webp",
    "gif":  "image/gif",
}

# Kimi vision OCR 的 system prompt
_VISION_SYSTEM = (
    "你是一名专业的幼儿园课程助手，擅长识别和理解幼师手写或打印的教学文档。"
    "请仔细阅读图片中的所有文字内容，原文提取，不要遗漏任何表格、标题、正文。"
    "如有表格，请用「字段名：内容」的格式逐行列出，保留结构。"
)

_VISION_USER = (
    "请识别并提取图片中的所有文字内容。"
    "要求：\n"
    "1. 完整提取，不遗漏任何文字\n"
    "2. 表格内容按「字段：内容」格式列出\n"
    "3. 标题单独一行\n"
    "4. 直接输出文字内容，不需要额外解释"
)

# Vision 模型：8k 上下文够一张周计划图，最快最省
_VISION_MODEL = "moonshot-v1-8k-vision-preview"


def _extract_image(file_bytes: bytes, ctx: DocContext) -> None:
    """
    通过 Kimi Vision OCR 识别图片中的文字。
    将识别结果写入 ctx.raw_text。
    """
    from core.settings import DASHSCOPE_API_KEY, DASHSCOPE_BASE_URL
    from openai import OpenAI

    if not DASHSCOPE_API_KEY:
        raise RuntimeError("未配置 DASHSCOPE_API_KEY，无法使用 Kimi Vision OCR")

    ext = ctx.filename.lower().rsplit(".", 1)[-1]
    mime = _IMAGE_MIME.get(ext, "image/jpeg")
    b64 = base64.b64encode(file_bytes).decode("ascii")
    data_url = f"data:{mime};base64,{b64}"

    # Kimi 用 Moonshot base_url
    vision_client = OpenAI(
        api_key=DASHSCOPE_API_KEY,
        base_url=DASHSCOPE_BASE_URL,
    )

    resp = vision_client.chat.completions.create(
        model=_VISION_MODEL,
        messages=[
            {"role": "system", "content": _VISION_SYSTEM},
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {"url": data_url},
                    },
                    {
                        "type": "text",
                        "text": _VISION_USER,
                    },
                ],
            },
        ],
        temperature=0,
        max_tokens=2000,
    )

    ocr_text = resp.choices[0].message.content.strip()
    logger.info("Kimi Vision OCR 完成 [%s]，识别字数：%d", ctx.filename, len(ocr_text))
    ctx.raw_text = ocr_text


# ──────────────────────────────────────────────
# 工具函数
# ──────────────────────────────────────────────

def _clean_cell(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def _dedup_adjacent(items: list[str]) -> list[str]:
    """去掉相邻重复项（合并单元格导致的）。"""
    result: list[str] = []
    for item in items:
        if not result or item != result[-1]:
            result.append(item)
    return result
