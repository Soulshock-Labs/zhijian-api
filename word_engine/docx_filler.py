"""
word_engine.docx_filler — python-docx 核心填充引擎 ⭐

铁律：遍历 Word 表格 → 关键字匹配单元格 → 保留原样式精准填充，格式分毫不动。

对外接口：
  fill_word_template()              — 路由层调用入口（自动选引擎）
  _fill_word_template_docx_bytes()  — 纯 docx 回填（可靠回退）
  docx_to_pdf_bytes()               — DOCX → PDF（需 Aspose）
  docx_to_images_bytes()            — DOCX → 图片列表（需 Aspose + pdf2image）
"""
from __future__ import annotations

import io
import logging
import os
import re
import subprocess
import tempfile
from typing import Optional
from urllib.parse import quote

from docx import Document
from docx.oxml.ns import qn
from fastapi.responses import StreamingResponse

from .field_map import (
    match_field,
    _ACTIVITY_FIELDS,
    ACTIVITY_LABEL_MAP,
    WEEKDAY_TAGS,
    _weekday_tag_from_header,
    _build_weekday_domain_plan,
)
from .template_tools import clean_template_keep_style, _is_colored_cell, _today_str

logger = logging.getLogger(__name__)

APP_VERSION = os.getenv("APP_VERSION", "1.2.1")
ENABLE_ASPOSE_WORDS = str(os.getenv("ENABLE_ASPOSE_WORDS", "0")).strip().lower() in {"1", "true", "yes", "on"}


# ──────────────────────────────────────────────
# docx 工具集
# ──────────────────────────────────────────────

def _get_cell_text(cell) -> str:
    return "".join(p.text for p in cell.paragraphs).strip()


def _copy_run_format(src_run, dst_run) -> None:
    """将 src_run 的字体/颜色格式复制到 dst_run。"""
    if src_run.font.name:
        dst_run.font.name = src_run.font.name
        rPr = dst_run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), src_run.font.name)
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.bold is not None:
        dst_run.font.bold = src_run.font.bold
    if src_run.font.color and src_run.font.color.type is not None:
        try:
            dst_run.font.color.rgb = src_run.font.color.rgb
        except Exception:
            pass


def _get_day_col_map(table) -> tuple[int, list[int]]:
    """扫描表格找到星期表头行及其列索引。返回 (row_idx, [col_indices])；未找到返回 (-1, [])。"""
    weekdays = ("星期一", "星期二", "星期三", "星期四", "星期五", "周一", "周二", "周三", "周四", "周五")
    for ri, row in enumerate(table.rows):
        cells = row.cells
        day_cols = [ci for ci, c in enumerate(cells) if any(d in "".join(_get_cell_text(c).split()) for d in weekdays)]
        if len(day_cols) >= 3:
            return ri, sorted(day_cols)
    return -1, []


def _write_cell_preserve_style(cell, text: str) -> None:
    """向已净空的单元格写入文本，保留第一个 run 的字体样式。"""
    if not cell.paragraphs:
        for line in text.split("\n"):
            cell.add_paragraph(line)
        return
    lines = text.split("\n")
    template_run = None
    for para in cell.paragraphs:
        for run in para.runs:
            template_run = run
            break
        if template_run:
            break
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    for i, line in enumerate(lines):
        if i < len(cell.paragraphs):
            para = cell.paragraphs[i]
            if para.runs:
                para.runs[0].text = line
                if template_run and para.runs[0] is not template_run:
                    _copy_run_format(template_run, para.runs[0])
            else:
                r = para.add_run(line)
                if template_run:
                    _copy_run_format(template_run, r)
        else:
            new_p = cell.add_paragraph()
            r = new_p.add_run(line)
            if template_run:
                _copy_run_format(template_run, r)


# ──────────────────────────────────────────────
# 周计划填充字段构建
# ──────────────────────────────────────────────

def _build_weekly_fill_data(
    theme: str,
    phil: str,
    ai_content: dict,
    activities: list[str],
    child_initiative: bool,
    child_desc: str,
    *,
    class_level: str = "",
    fill_unselected: bool = False,
) -> dict[str, str]:
    """周计划填充字段字典（Aspose / python-docx 共用）。"""
    today = _today_str()
    goals_text = "\n".join(f"目标{i+1}：{g}" for i, g in enumerate(ai_content.get("goals", [])))
    prep_text = "\n".join(f"• {p}" for p in ai_content.get("preparation", []))
    weekly_targets = ai_content.get("weekly_targets", {}) if isinstance(ai_content, dict) else {}
    weekly_teaching = str((weekly_targets or {}).get("teaching", "") or "").strip()
    weekly_life = str((weekly_targets or {}).get("life", "") or "").strip()
    weekly_family = str((weekly_targets or {}).get("family", "") or "").strip()
    weekly_environment = str((weekly_targets or {}).get("environment", "") or "").strip()

    fill_data: dict[str, str] = {
        "theme":             theme,
        "philosophy":        phil,
        "goals":             goals_text,
        "preparation":       prep_text,
        "evaluation":        ai_content.get("evaluation", ""),
        "week_overview":     "、".join([ACTIVITY_LABEL_MAP.get(a, a) for a in activities]) or "本周按班级节奏灵活安排",
        "daily_points":      "",
        "key_difficulty":    weekly_teaching or "重点：围绕主题形成连续经验；难点：兼顾个体差异与活动节奏。",
        "children_baseline": child_desc if child_desc else "基于前期观察记录，幼儿对本周主题已有初步兴趣与经验基础。",
        "resource_list":     prep_text,
        "observation_plan":  ai_content.get("evaluation", ""),
        "individual_support": ai_content.get("child_initiative_note", "") or "关注不同发展水平幼儿，提供分层支持与差异化引导。",
        "safety_risk":       "重点关注户外活动与材料使用安全；提前进行规则提醒与风险巡视。",
        "next_week_plan":    "根据本周观察结果调整下周材料投放与活动难度，延续幼儿高兴趣点。",
        "nap_guidance":      ai_content.get("nap_guidance", ""),
        "class_info":        f"________班    日期：{today}    天气：☀️ 晴",
        "child_initiative": (
            f"✅ 本周有幼儿自主发起活动\n{child_desc}\n"
            f"💡 {phil}理念中，幼儿自主发起的活动是最宝贵的课程生长点，请及时记录跟进。"
            if child_initiative and child_desc
            else ("✅ 本周有幼儿自主发起活动" if child_initiative else "")
        ),
    }
    if weekly_life:
        fill_data["life"] = weekly_life
    if weekly_family:
        fill_data["family"] = weekly_family
    if weekly_environment:
        fill_data["environment"] = weekly_environment

    weekday_domain = _build_weekday_domain_plan(theme)
    day_zh = {"mon": "周一", "tue": "周二", "wed": "周三", "thu": "周四", "fri": "周五"}
    fill_data["daily_points"] = "\n".join(
        f"{day_zh[tag]}（{weekday_domain.get(tag, '综合')}）：围绕「{theme}」推进核心经验，体现{phil}特色。"
        for tag in ("mon", "tue", "wed", "thu", "fri")
    )

    ai_acts: dict = ai_content.get("activities", {})
    for act_id in ACTIVITY_LABEL_MAP:
        content = ai_acts.get(act_id, "")
        if act_id in activities and content:
            fill_data[act_id] = content
        elif fill_unselected:
            fill_data[act_id] = "（本周未启用该板块，可按班级实际勾选后再生成）"
        base_for_day = fill_data.get(act_id, "") or content
        if base_for_day:
            for tag in ("mon", "tue", "wed", "thu", "fri"):
                domain = weekday_domain.get(tag, "综合")
                fill_data[f"{act_id}__{tag}"] = f"{base_for_day}\n【{day_zh[tag]}·{domain}】结合班级现状分层引导。"

    guidance_items = ai_content.get("guidance", [])
    if guidance_items:
        fill_data["guidance"] = "\n".join(f"{i+1}. {g}" for i, g in enumerate(guidance_items))

    _day_prefix_re = re.compile(r"^(?:周[一二三四五]|星期[一二三四五])[^\S\n]*[：:·\-\s]*")
    tag_order = ("mon", "tue", "wed", "thu", "fri")
    is_xiaopan = "小班" in class_level
    for field_id, field_default in (
        ("study", "本周集中活动（按班级课程表安排）"),
        ("game",  fill_data.get("area", "")),
    ):
        if field_id in fill_data:
            continue
        raw_text = ai_acts.get(field_id, "").strip()
        if not raw_text:
            fill_data[field_id] = field_default
            continue
        lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
        per_day = [_day_prefix_re.sub("", l).strip() for l in lines]
        if len(per_day) >= 2:
            fill_data[field_id] = per_day[0]
            for i, tag in enumerate(tag_order):
                day_content = per_day[i] if i < len(per_day) else per_day[-1]
                if is_xiaopan and field_id == "study" and "；" in day_content:
                    day_content = day_content.split("；")[0].strip()
                fill_data[f"{field_id}__{tag}"] = day_content
        else:
            fill_data[field_id] = raw_text
    return fill_data


# ──────────────────────────────────────────────
# 核心填充函数
# ──────────────────────────────────────────────

def _fill_word_template_docx_bytes(cleaned_bytes: bytes, fill_data: dict[str, str]) -> bytes:
    """python-docx 回填（Aspose 不可用或抛错时的可靠回退）。"""
    doc = Document(io.BytesIO(cleaned_bytes))
    filled_tc_elems: set = set()

    for table in doc.tables:
        rows = table.rows
        day_header_ri, day_cols = _get_day_col_map(table)

        if day_header_ri >= 0 and len(day_cols) >= 3:
            day_header_tc_elems: set = set()
            hdr_row = rows[day_header_ri]
            for day_ci in day_cols:
                if day_ci < len(hdr_row.cells):
                    day_header_tc_elems.add(hdr_row.cells[day_ci]._tc)
            header_cache: dict[int, str] = {}
            for day_ci in day_cols:
                if day_ci < len(hdr_row.cells):
                    header_cache[day_ci] = _get_cell_text(hdr_row.cells[day_ci])

            for row_idx in range(day_header_ri):
                cells = rows[row_idx].cells
                for col_idx, cell in enumerate(cells):
                    field = match_field(_get_cell_text(cell))
                    content = fill_data.get(field) if field else None
                    if not content:
                        continue
                    for ci in range(col_idx + 1, len(cells)):
                        cand = cells[ci]
                        if match_field(_get_cell_text(cand)) or _is_colored_cell(cand):
                            continue
                        tc_elem = cand._tc
                        if tc_elem in day_header_tc_elems:
                            continue
                        if tc_elem in filled_tc_elems:
                            break
                        filled_tc_elems.add(tc_elem)
                        _write_cell_preserve_style(cand, content)
                        break

            hdr_row = rows[day_header_ri]
            for day_ci, txt in header_cache.items():
                if day_ci < len(hdr_row.cells):
                    _write_cell_preserve_style(hdr_row.cells[day_ci], txt)

            current_field: str | None = None
            for row_idx in range(day_header_ri + 1, len(rows)):
                cells = rows[row_idx].cells
                if not cells:
                    continue
                col0_field = match_field(_get_cell_text(cells[0])) if _get_cell_text(cells[0]) else None
                if col0_field:
                    current_field = col0_field
                row_field = current_field
                if len(cells) > 1:
                    col1_text = _get_cell_text(cells[1])
                    col1_field = match_field(col1_text) if col1_text else None
                    if col1_field and col1_field in _ACTIVITY_FIELDS:
                        row_field = col1_field
                if not row_field:
                    continue
                base_content = fill_data.get(row_field)
                day_tc_ids = {id(cells[ci]._tc) for ci in day_cols if ci < len(cells)}
                all_days_merged = len(day_tc_ids) == 1
                for day_ci in day_cols:
                    if day_ci >= len(cells):
                        continue
                    target = cells[day_ci]
                    if _is_colored_cell(target):
                        continue
                    day_tag = _weekday_tag_from_header(header_cache.get(day_ci, ""))
                    if all_days_merged or not day_tag:
                        cell_content = base_content or fill_data.get(f"{row_field}__mon")
                    else:
                        cell_content = fill_data.get(f"{row_field}__{day_tag}") or base_content
                    if not cell_content:
                        continue
                    tc_elem = target._tc
                    if tc_elem in filled_tc_elems:
                        continue
                    filled_tc_elems.add(tc_elem)
                    _write_cell_preserve_style(target, cell_content)
        else:
            for row_idx, row in enumerate(rows):
                cells = row.cells
                for col_idx, cell in enumerate(cells):
                    field = match_field(_get_cell_text(cell))
                    content = fill_data.get(field) if field else None
                    if not content:
                        continue
                    target = None
                    if col_idx + 1 < len(cells):
                        cand = cells[col_idx + 1]
                        if not match_field(_get_cell_text(cand)) and not _is_colored_cell(cand):
                            target = cand
                    if target is None and row_idx + 1 < len(rows):
                        cand = rows[row_idx + 1].cells[col_idx]
                        if not match_field(_get_cell_text(cand)) and not _is_colored_cell(cand):
                            target = cand
                    if target is None:
                        continue
                    tc_elem = target._tc
                    if tc_elem in filled_tc_elems:
                        continue
                    filled_tc_elems.add(tc_elem)
                    _write_cell_preserve_style(target, content)

    try:
        doc.core_properties.comments = (
            f"小纸笺导出 · v{APP_VERSION} · 排版引擎 python-docx（Aspose 回退）"
        )
    except Exception:
        pass
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _build_content_disposition(filename: str) -> str:
    """构造兼容中英文文件名的 Content-Disposition 下载头。"""
    filename = filename.strip() or "export.docx"
    utf8_name = quote(filename)
    ascii_fallback = re.sub(r"[^A-Za-z0-9._-]+", "_", filename) or "export.docx"
    return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{utf8_name}"


def fill_word_template(
    template_bytes: bytes,
    theme: str,
    phil: str,
    ai_content: dict,
    activities: list[str],
    child_initiative: bool,
    child_desc: str,
    *,
    class_level: str = "",
    fill_unselected: bool = False,
) -> tuple[bytes, str]:
    """
    铁律填充入口：净空模板骨架 → 写入 AI 内容。
    优先 Aspose.Words；失败时自动回退 python-docx。
    返回 (docx_bytes, engine_name)。
    """
    cleaned_bytes = clean_template_keep_style(template_bytes)
    fill_data = _build_weekly_fill_data(
        theme, phil, ai_content, activities, child_initiative, child_desc,
        class_level=class_level, fill_unselected=fill_unselected,
    )
    if ENABLE_ASPOSE_WORDS:
        try:
            from .aspose_filler import _fill_word_template_aspose_bytes
            return _fill_word_template_aspose_bytes(cleaned_bytes, fill_data), "aspose-words"
        except Exception as e:
            logger.warning("Aspose 周计划填充失败，回退 python-docx：%s", e, exc_info=True)
    return _fill_word_template_docx_bytes(cleaned_bytes, fill_data), "python-docx"


# ──────────────────────────────────────────────
# 格式转换
# ──────────────────────────────────────────────

def docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """DOCX → PDF（需 Aspose，ENABLE_ASPOSE_WORDS=1）。"""
    from .aspose_filler import _aw_require, _aw_doc_from_bytes, aw as _aw
    _aw_require()
    doc = _aw_doc_from_bytes(docx_bytes)
    out = io.BytesIO()
    doc.save(out, _aw.SaveFormat.PDF)
    out.seek(0)
    return out.read()


def docx_to_images_bytes(docx_bytes: bytes, format: str = "png", dpi: int = 150) -> list[bytes]:
    """DOCX → PDF → 图片列表（每页一张）。format: 'png' 或 'jpg'。"""
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        raise RuntimeError("pdf2image 未安装，无法导出图片。请运行 pip install pdf2image")

    try:
        pdf_bytes = docx_to_pdf_bytes(docx_bytes)
    except Exception:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(docx_bytes)
            docx_path = f.name
        try:
            pdf_path = docx_path.replace(".docx", ".pdf")
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", docx_path],
                check=True, capture_output=True, timeout=30,
            )
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
        finally:
            import os as _os
            _os.unlink(docx_path)
            if _os.path.exists(pdf_path):
                _os.unlink(pdf_path)

    images = convert_from_bytes(pdf_bytes, dpi=dpi, fmt=format)
    out_bytes = []
    for img in images:
        buf = io.BytesIO()
        img.save(buf, format=format.upper() if format != "jpg" else "JPEG")
        buf.seek(0)
        out_bytes.append(buf.read())
    return out_bytes
