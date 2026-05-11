"""
word_engine.template_tools — 模板净空 / 分析 / 内置模板构建

职责：
  clean_template_keep_style() — 净空模板，保留所有样式
  analyze_template_docx()     — 模板结构自检，返回合规报告
  _build_standard_weekly_template_bytes() — 内置标准周模板
  _build_standard_daily_template_bytes()  — 内置标准日教案模板
  _is_colored_cell()          — 着色单元格检测（固定区保护）
"""
from __future__ import annotations

import io
import logging
import re
from datetime import date
from typing import Optional

from docx import Document
from docx.oxml.ns import qn

from kindergarten_template_cleaner import clean_docx_bytes, find_all_day_header_rows
from .field_map import match_field, WEEKDAY_TAGS

logger = logging.getLogger(__name__)

WEEKDAY_HEADERS = (
    "星期一", "星期二", "星期三", "星期四", "星期五",
    "周一", "周二", "周三", "周四", "周五",
)
ROW_LABEL_HINTS = (
    "活动主题", "教学主题", "活动目标", "活动准备", "生活活动", "学习活动",
    "游戏活动", "区域活动", "户外活动", "家园活动", "离园活动", "评价与反思", "反思",
)

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")


def _today_str() -> str:
    d = date.today()
    return f"{d.year}年{d.month}月{d.day}日"


def _get_cell_text(cell) -> str:
    return "".join(p.text for p in cell.paragraphs).strip()


def _is_colored_cell(cell) -> bool:
    """判断单元格是否带底色（带底色视为模板固定区，不可动）。"""
    try:
        tc_pr = cell._tc.tcPr
        if tc_pr is None:
            return False
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            return False
        fill = (shd.get(qn("w:fill")) or "").strip().upper()
        if not fill or fill in {"AUTO", "FFFFFF"}:
            return False
        return True
    except Exception:
        return False


def clean_template_keep_style(template_bytes: bytes) -> bytes:
    """净空模板：委托 kindergarten_template_cleaner，失败时回退原始字节。"""
    try:
        return clean_docx_bytes(template_bytes)
    except Exception as e:
        logger.warning("净空模板失败，使用原始 .docx：%s", e, exc_info=True)
        return template_bytes


def _build_standard_weekly_template_bytes() -> bytes:
    """生成内置标准周模板（19 模块超集，含可识别标签）。"""
    doc = Document()
    doc.add_heading("幼儿园周计划（标准模板 · 19模块）", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    fields = [
        ("基础信息", "园所：________    班级：________    周次：________    日期：________"),
        ("本周主题", ""),
        ("教育理念", ""),
        ("周总目标（五大领域）", ""),
        ("本周重点与难点", ""),
        ("幼儿已有经验", ""),
        ("本周活动总览表", ""),
        ("每日活动要点", ""),
        ("区域活动设计", ""),
        ("户外与体能活动", ""),
        ("生活活动与保育", ""),
        ("环境创设", ""),
        ("家园共育", ""),
        ("个别化支持", ""),
        ("安全与风险提示", ""),
        ("资源与材料清单", ""),
        ("观察记录计划", ""),
        ("周反思", ""),
        ("下周衔接", ""),
        ("午睡指导", ""),
        ("教师签名", "主班：________    配班：________"),
    ]
    for left, right in fields:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _build_standard_daily_template_bytes() -> bytes:
    """生成内置标准日教案空白模板（含可识别标签）。"""
    doc = Document()
    doc.add_heading("幼儿园日教案（标准模板）", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    fields = [
        ("活动名称", ""),
        ("教育理念", ""),
        ("班级", "________班    日期：________"),
        ("活动目标", ""),
        ("活动准备", ""),
        ("活动导入", ""),
        ("活动过程", ""),
        ("活动延伸", ""),
        ("活动反思", ""),
        ("观察要点", ""),
    ]
    for left, right in fields:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def analyze_template_docx(template_bytes: bytes) -> dict:
    """对上传模板做结构级自检，输出与 TEMPLATE_STANDARD.md v1.1 对齐的报告。"""
    from .field_map import match_field as _match_field  # local alias to avoid shadowing

    doc = Document(io.BytesIO(template_bytes))
    day_header_rows = find_all_day_header_rows(template_bytes)
    table_stats: list[dict] = []
    all_text_joined: list[str] = []
    colored_cells = 0
    placeholder_tags: list[str] = []
    weekday_row_hits = 0
    field_hits: dict[str, int] = {}
    daily_hits: dict[str, int] = {}
    nested_label_rows = 0
    outdoor_nested_hint = False

    # 日教案字段匹配（延迟导入避免循环）
    try:
        from word_engine.daily_field_map import match_daily_field
    except ImportError:
        def match_daily_field(_t: str):  # type: ignore[misc]
            return None

    for ti, table in enumerate(doc.tables):
        rows = table.rows
        max_cols = max((len(r.cells) for r in rows), default=0)
        dhr = day_header_rows[ti] if ti < len(day_header_rows) else -1
        table_stats.append({"table_index": ti + 1, "rows": len(rows), "max_cols": max_cols, "day_header_row": dhr})

        for row in rows:
            cells = row.cells
            first_text = _get_cell_text(cells[0]) if cells else ""
            if first_text and any(h in first_text for h in ROW_LABEL_HINTS) and len(cells) >= 5:
                nested_label_rows += 1
            for cell in cells:
                t = _get_cell_text(cell)
                if t:
                    all_text_joined.append(t)
                if _is_colored_cell(cell):
                    colored_cells += 1
                for m in _PLACEHOLDER_RE.finditer(t):
                    placeholder_tags.append(m.group(1).strip())
                for day in WEEKDAY_HEADERS:
                    if day in t:
                        weekday_row_hits += 1
                        break
                mf = _match_field(t)
                if mf:
                    field_hits[mf] = field_hits.get(mf, 0) + 1
                md = match_daily_field(t)
                if md:
                    daily_hits[md] = daily_hits.get(md, 0) + 1

    blob = "".join(all_text_joined)
    if "户外" in blob and ("自主" in blob or "集体" in blob):
        outdoor_nested_hint = True

    uniq_ph = list(dict.fromkeys(placeholder_tags))
    if uniq_ph:
        doc_type = "placeholder_template"
    elif len([k for k in daily_hits if k in ("introduction", "process", "extension", "reflection")]) >= 2:
        doc_type = "daily_plan"
    elif ("星期一" in blob and "星期五" in blob) or ("周一" in blob and "周五" in blob) or weekday_row_hits >= 8:
        doc_type = "weekly_grid"
    elif field_hits:
        doc_type = "general_plan"
    else:
        doc_type = "unknown"

    mappings: list[dict] = []
    if outdoor_nested_hint:
        subs = [s for s in ("自主", "集体") if s in blob]
        mappings.append({"tag": "plan.outdoor", "status": "partial_match",
                         "reason": f"nested sub-rows hints: {subs}" if subs else "outdoor section present"})

    score = 0.5
    if doc_type == "weekly_grid":
        score += 0.18
    elif doc_type == "daily_plan":
        score += 0.2
    elif doc_type == "placeholder_template":
        score += 0.22
    elif doc_type == "general_plan":
        score += 0.12
    if field_hits or daily_hits:
        score += 0.1
    if nested_label_rows:
        score -= 0.04
    score = round(min(0.95, max(0.4, score)), 2)

    return {
        "status": "success",
        "standard_ref": "TEMPLATE_STANDARD.md v1.2.0",
        "confidence_score": score,
        "doc_type_guess": doc_type,
        "document": {
            "tables": len(doc.tables),
            "day_header_rows": day_header_rows,
            "table_stats": table_stats,
            "colored_cells_detected": colored_cells,
            "placeholder_count": len(uniq_ph),
            "placeholders_sample": uniq_ph[:24],
            "weekday_cell_scan_hits": weekday_row_hits,
            "nested_wide_rows_hint": nested_label_rows,
        },
        "keyword_hits": {"fill_fields": field_hits, "daily_fields": daily_hits},
        "mappings": mappings,
        "compliance": {
            "checklist": [
                "星期表头行（动态）与首列在净空引擎中默认保护；非星期表头的第一行按普通行处理",
                "下载「原文件名_净空模板.docx」与原文件并排对比表格是否变形",
                "若含 {{占位符}}，回填优先绑定占位符槽位；Word 可能拆分 run，检测需拼合文本",
            ],
        },
    }
